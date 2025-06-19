import pandas as pd
import json
import os
from typing import Dict, Any, List
import uuid
from sqlalchemy.orm import Session
from models import UploadedFile, SemanticModel
import PyPDF2
import docx
import numpy as np
import time
import asyncio

class FileProcessorService:
    def __init__(self):
        self.supported_types = {
            'csv': self._process_csv,
            'xlsx': self._process_excel,
            'xls': self._process_excel,
            'json': self._process_json,
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'txt': self._process_text
        }
    
    def _update_file_status(self, file_id: str, status: str, extracted_data: Dict = None, metadata: Dict = None, error: str = None):
        """Update file status in database"""
        from database import SessionLocal
        
        db = SessionLocal()
        try:
            file_record = db.query(UploadedFile).filter(UploadedFile.id == file_id).first()
            if file_record:
                file_record.processing_status = status
                if extracted_data:
                    file_record.extracted_data = extracted_data
                if metadata:
                    file_record.metadata = metadata
                if error:
                    file_record.metadata = {"error": error}
                
                db.commit()
                print(f"Updated file {file_id} status to: {status}")
        except Exception as e:
            print(f"Error updating file status: {str(e)}")
            db.rollback()
        finally:
            db.close()

    async def process_file(self, file_id: str, file_path: str, file_type: str):
        """Process uploaded file and extract data"""
        
        try:
            print(f"Starting processing for file {file_id}, type: {file_type}")
            
            # Update status to processing with start time
            start_metadata = {"started_at": time.time()}
            self._update_file_status(file_id, "processing", metadata=start_metadata)
            
            # Small delay to show processing status
            await asyncio.sleep(1)
            
            # Process file based on type
            processor = self.supported_types.get(file_type.lower())
            if not processor:
                self._update_file_status(file_id, "failed", error="Unsupported file type")
                return
            
            print(f"Processing file with processor for type: {file_type}")
            # Run the synchronous processor in a thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            extracted_data = await loop.run_in_executor(None, processor, file_path)
            print(f"Successfully extracted data, rows: {len(extracted_data.get('rows', []))}")
            
            # Update file record with extracted data
            metadata = {
                "rows": len(extracted_data.get('rows', [])),
                "columns": len(extracted_data.get('columns', [])),
                "data_types": extracted_data.get('data_types', {}),
                "file_info": extracted_data.get('file_info', {}),
                "started_at": start_metadata["started_at"],
                "completed_at": time.time()
            }
            
            self._update_file_status(file_id, "completed", extracted_data, metadata)
            print(f"File {file_id} processing completed successfully")
            
        except Exception as e:
            print(f"Error processing file {file_id}: {str(e)}")
            self._update_file_status(file_id, "failed", error=str(e))
    
    def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Process CSV file"""
        try:
            print(f"Processing CSV file: {file_path}")
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    print(f"Successfully read CSV with encoding: {encoding}")
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                raise ValueError("Could not read CSV file with any supported encoding")
            
            result = self._dataframe_to_dict(df)
            print(f"CSV processing complete. Rows: {len(result.get('rows', []))}, Columns: {len(result.get('columns', []))}")
            return result
            
        except Exception as e:
            print(f"Error processing CSV: {str(e)}")
            raise Exception(f"Error processing CSV: {str(e)}")
    
    def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Process Excel file"""
        try:
            print(f"Processing Excel file: {file_path}")
            # Read first sheet
            df = pd.read_excel(file_path, sheet_name=0)
            result = self._dataframe_to_dict(df)
            print(f"Excel processing complete. Rows: {len(result.get('rows', []))}, Columns: {len(result.get('columns', []))}")
            return result
            
        except Exception as e:
            print(f"Error processing Excel: {str(e)}")
            raise Exception(f"Error processing Excel: {str(e)}")
    
    def _process_json(self, file_path: str) -> Dict[str, Any]:
        """Process JSON file"""
        try:
            print(f"Processing JSON file: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # If data is a list of objects, convert to DataFrame
            if isinstance(data, list) and all(isinstance(item, dict) for item in data):
                df = pd.DataFrame(data)
                return self._dataframe_to_dict(df)
            
            # If data is a single object, treat as one row
            elif isinstance(data, dict):
                df = pd.DataFrame([data])
                return self._dataframe_to_dict(df)
            
            # Otherwise, return as text content
            else:
                return {
                    "type": "text",
                    "content": json.dumps(data, indent=2),
                    "file_info": {"original_type": "json"}
                }
                
        except Exception as e:
            print(f"Error processing JSON: {str(e)}")
            raise Exception(f"Error processing JSON: {str(e)}")
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF file"""
        try:
            print(f"Processing PDF file: {file_path}")
            text_content = ""
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
            
            result = {
                "type": "text",
                "content": text_content,
                "file_info": {
                    "pages": len(pdf_reader.pages),
                    "original_type": "pdf"
                }
            }
            print(f"PDF processing complete. Pages: {len(pdf_reader.pages)}")
            return result
            
        except Exception as e:
            print(f"Error processing PDF: {str(e)}")
            raise Exception(f"Error processing PDF: {str(e)}")
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process Word document"""
        try:
            print(f"Processing DOCX file: {file_path}")
            doc = docx.Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            result = {
                "type": "text",
                "content": text_content,
                "file_info": {
                    "paragraphs": len(doc.paragraphs),
                    "original_type": "docx"
                }
            }
            print(f"DOCX processing complete. Paragraphs: {len(doc.paragraphs)}")
            return result
            
        except Exception as e:
            print(f"Error processing DOCX: {str(e)}")
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process text file"""
        try:
            print(f"Processing text file: {file_path}")
            encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    print(f"Successfully read text file with encoding: {encoding}")
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("Could not read text file with any supported encoding")
            
            result = {
                "type": "text",
                "content": content,
                "file_info": {
                    "length": len(content),
                    "lines": len(content.split('\n')),
                    "original_type": "txt"
                }
            }
            print(f"Text processing complete. Length: {len(content)} chars")
            return result
            
        except Exception as e:
            print(f"Error processing text file: {str(e)}")
            raise Exception(f"Error processing text file: {str(e)}")
    
    def _dataframe_to_dict(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Convert DataFrame to dictionary format"""
        print(f"Converting DataFrame to dict. Shape: {df.shape}")
        
        # Clean the dataframe
        df = self._clean_dataframe(df)
        
        # Get column info
        columns = []
        data_types = {}
        
        for col in df.columns:
            col_type = self._infer_column_type(df[col])
            columns.append({
                "name": str(col),
                "type": col_type
            })
            data_types[str(col)] = col_type
        
        # Convert to records
        rows = df.to_dict('records')
        
        # Handle NaN values
        for row in rows:
            for key, value in row.items():
                if pd.isna(value):
                    row[key] = None
                elif isinstance(value, (np.int64, np.int32)):
                    row[key] = int(value)
                elif isinstance(value, (np.float64, np.float32)):
                    row[key] = float(value)
        
        result = {
            "type": "tabular",
            "columns": columns,
            "rows": rows,
            "row_count": len(df),
            "column_count": len(df.columns),
            "data_types": data_types,
            "file_info": {
                "shape": df.shape,
                "memory_usage": df.memory_usage(deep=True).sum()
            }
        }
        
        print(f"DataFrame conversion complete. Final rows: {len(rows)}, columns: {len(columns)}")
        return result
    
    def _clean_dataframe(self, df: pd.DataFrame) -> pd.DataFrame:
        """Clean and prepare dataframe"""
        print(f"Cleaning DataFrame. Original shape: {df.shape}")
        
        # Remove completely empty rows and columns
        df = df.dropna(how='all').dropna(axis=1, how='all')
        
        # Convert column names to strings
        df.columns = [str(col) for col in df.columns]
        
        # Handle duplicate column names
        df.columns = pd.io.common.dedup_names(df.columns, is_potential_multiindex=False)
        
        print(f"DataFrame cleaned. New shape: {df.shape}")
        return df
    
    def _infer_column_type(self, series: pd.Series) -> str:
        """Infer the type of a pandas Series"""
        if pd.api.types.is_numeric_dtype(series):
            if pd.api.types.is_integer_dtype(series):
                return "integer"
            else:
                return "number"
        elif pd.api.types.is_datetime64_any_dtype(series):
            return "datetime"
        elif pd.api.types.is_bool_dtype(series):
            return "boolean"
        else:
            return "string"
    
    async def create_semantic_model(self, extracted_data: Dict[str, Any], model_name: str) -> SemanticModel:
        """Create a semantic model from extracted data"""
        from database import SessionLocal
        
        db = SessionLocal()
        try:
            if extracted_data.get("type") != "tabular":
                raise ValueError("Can only create semantic models from tabular data")
            
            # Generate schema definition
            schema_definition = {
                "tables": {
                    "main_table": {
                        "sql": "SELECT * FROM uploaded_data",
                        "columns": {}
                    }
                },
                "metrics": [],
                "dimensions": [],
                "joins": []
            }
            
            # Add columns
            for col in extracted_data.get("columns", []):
                col_name = col["name"]
                col_type = col["type"]
                
                schema_definition["tables"]["main_table"]["columns"][col_name] = {
                    "type": col_type,
                    "primary": False
                }
                
                # Add as dimension or metric
                if col_type in ["string", "boolean"]:
                    schema_definition["dimensions"].append({
                        "name": col_name.lower().replace(" ", "_"),
                        "title": col_name,
                        "sql": f"main_table.{col_name}",
                        "type": col_type
                    })
                elif col_type in ["number", "integer"]:
                    # Add as both metric and dimension
                    schema_definition["metrics"].append({
                        "name": f"sum_{col_name.lower().replace(' ', '_')}",
                        "title": f"Total {col_name}",
                        "type": "sum",
                        "sql": f"SUM(main_table.{col_name})",
                        "format": "number"
                    })
                    
                    schema_definition["dimensions"].append({
                        "name": col_name.lower().replace(" ", "_"),
                        "title": col_name,
                        "sql": f"main_table.{col_name}",
                        "type": col_type
                    })
            
            # Create semantic model
            semantic_model = SemanticModel(
                name=model_name,
                description=f"Auto-generated semantic model for {model_name}",
                schema_definition=schema_definition
            )
            
            db.add(semantic_model)
            db.commit()
            db.refresh(semantic_model)
            
            print(f"Created semantic model: {model_name}")
            return semantic_model
            
        finally:
            db.close()